"""Format-aware text extraction for grounding sources.

Each grounding source goes through ``extract_text`` which returns
``(text, kind)``. The kind label tells ``_read_sources`` what fallback
chain to apply:

- ``text``       - already plain text (.txt, .md, .rst); read as UTF-8
- ``pdf-text``   - PDF with extractable text streams
- ``pdf-scanned``- PDF that yielded near-empty text (image-only / scan)
- ``docx``       - Word .docx (extracted via python-docx)
- ``odt``        - OpenDocument .odt (extracted via odfpy)
- ``rtf``        - .rtf (extracted via striprtf)
- ``html``       - .html / .htm (extracted via stdlib html.parser)

Scanned PDFs are not failures; the caller decides whether to look for a
sibling, run OCR, or skip the source. ``find_sibling`` resolves the
priority lookup. ``infer_pdf_language`` is a hint helper - it returns a
suggestion the agent should confirm; the OCR tool's ``lang=`` parameter
is the agent's responsibility to supply.
"""

from __future__ import annotations

from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

# ---------------------------------------------------------------------------
# Sibling lookup priority + excluded extensions
# ---------------------------------------------------------------------------

# First match wins. All paths same-stem in the same directory.
SIBLING_PRIORITY: tuple[str, ...] = (
    ".ocr.txt",
    ".txt",
    ".docx",
    ".doc",
    ".odt",
    ".md",
    ".rst",
    ".html",
    ".htm",
    ".rtf",
)

# Extensions that are NOT candidates for sibling-text fallback (image
# formats and the original PDF itself).
EXCLUDED_SIBLING_EXTENSIONS: frozenset[str] = frozenset(
    {
        ".pdf",
        ".jpg",
        ".jpeg",
        ".png",
        ".tiff",
        ".tif",
        ".gif",
        ".bmp",
        ".webp",
        ".heic",
        ".heif",
        ".svg",
    }
)


# ---------------------------------------------------------------------------
# Errors + result types
# ---------------------------------------------------------------------------


class UnsupportedFormat(Exception):
    """Raised when no extractor matches a source path.

    The error message is suitable for surfacing to the agent (it names
    the file kind and points at relevant external tools).
    """


@dataclass(frozen=True)
class Extracted:
    text: str
    kind: str
    page_count: int | None = None  # only meaningful for ``pdf-*`` kinds


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

# Heuristic: a PDF is treated as "scanned" when extracted text averages
# fewer than this many characters per page. Configurable via the CLI
# ``--scanned-threshold`` flag; passed through to is_scanned_pdf().
DEFAULT_SCANNED_THRESHOLD: int = 100


def extract_text(path: Path, *, scanned_threshold: int = DEFAULT_SCANNED_THRESHOLD) -> Extracted:
    """Dispatch to the right extractor based on file extension.

    Returns an ``Extracted`` carrying the full text plus a ``kind`` tag
    that the caller uses to route fallbacks. Raises ``UnsupportedFormat``
    on file types we cannot read - the caller surfaces that to the agent
    via the ack-warning gate so the source is skipped, not silently
    accepted as empty text.
    """
    suffix = path.suffix.lower()
    name_lower = path.name.lower()

    if name_lower.endswith(".ocr.txt"):
        return Extracted(text=_read_text_strict(path), kind="text")

    if suffix in {".txt", ".md", ".rst"}:
        return Extracted(text=_read_text_strict(path), kind="text")

    if suffix == ".pdf":
        text, page_count = _extract_pdf(path)
        kind = "pdf-scanned" if is_scanned_pdf(text, page_count, scanned_threshold) else "pdf-text"
        return Extracted(text=text, kind=kind, page_count=page_count)

    if suffix == ".docx":
        return Extracted(text=_extract_docx(path), kind="docx")

    if suffix == ".odt":
        return Extracted(text=_extract_odt(path), kind="odt")

    if suffix == ".rtf":
        return Extracted(text=_extract_rtf(path), kind="rtf")

    if suffix in {".html", ".htm"}:
        return Extracted(text=_extract_html(path), kind="html")

    if suffix == ".doc":
        # Legacy binary Word - we don't bundle a reader. Surface as
        # unsupported so the agent converts upstream (libreoffice
        # --headless --convert-to docx) and reruns.
        raise UnsupportedFormat(
            f"{path.name} is legacy .doc (OLE2 binary). Convert to .docx "
            f"first: `libreoffice --headless --convert-to docx {path.name}`"
        )

    raise UnsupportedFormat(
        f"{path.name} has unsupported extension {suffix!r}. Supported: "
        f"{sorted({'.txt', '.md', '.rst', '.pdf', '.docx', '.odt', '.rtf', '.html', '.htm'})}"
    )


def find_sibling(path: Path) -> Path | None:
    """Return the highest-priority same-stem sibling, or None.

    Walks ``path.parent`` once. A "sibling" is any file with the same
    stem (case-sensitive) and an extension in ``SIBLING_PRIORITY``.
    Excluded extensions (PDFs, images) are ignored even if they exist.
    """
    if not path.exists():
        return None
    stem = path.stem
    parent = path.parent
    candidates: dict[str, Path] = {}
    for sibling in parent.iterdir():
        if not sibling.is_file():
            continue
        if sibling == path:
            continue
        if sibling.name.startswith(stem + "."):
            ext_part = sibling.name[len(stem) :].lower()  # leading dot included
            if ext_part in EXCLUDED_SIBLING_EXTENSIONS:
                continue
            candidates[ext_part] = sibling
    for ext in SIBLING_PRIORITY:
        if ext in candidates:
            return candidates[ext]
    return None


def is_scanned_pdf(text: str, page_count: int | None, threshold: int) -> bool:
    """Classify a PDF as scanned when its text density is too low.

    ``page_count`` may be 0 when pypdf could not parse pages at all - we
    treat that as scanned too (zero extractable signal).
    """
    if page_count is None or page_count <= 0:
        return True
    return len(text.strip()) / page_count < threshold


def infer_pdf_language(text: str, path: Path) -> str:
    """Suggest an OCR language code from sparse extracted text + filename.

    HINT only - the agent confirms the actual ``--ocr-lang`` value. We
    return a best-effort 3-letter Tesseract code (``eng``, ``deu`` etc).
    Default is ``eng`` when nothing parses. ``langdetect`` is preferred
    when installed; otherwise we use a tiny ASCII-frequency heuristic.
    """
    sample = text.strip()
    if not sample:
        # Look at filename tokens - "report-fr.pdf" hints French.
        name_hint = _language_from_filename(path.name.lower())
        if name_hint:
            return name_hint
        return "eng"
    try:
        from langdetect import detect  # type: ignore[import-not-found]

        iso2 = detect(sample[:2000])
    except Exception:
        return _language_from_filename(path.name.lower()) or "eng"
    return _ISO2_TO_TESSERACT.get(iso2, "eng")


# ---------------------------------------------------------------------------
# Format-specific extractors
# ---------------------------------------------------------------------------


def _read_text_strict(path: Path) -> str:
    """Read a file as strict UTF-8, surface failure as UnsupportedFormat.

    Lets the caller fall back to the SKIPPED gate warning rather than
    propagating UnicodeDecodeError. Renamed binaries (PNG saved as .txt)
    end up here.
    """
    try:
        return path.read_text(encoding="utf-8", errors="strict")
    except UnicodeDecodeError as exc:
        raise UnsupportedFormat(
            f"{path.name} is not valid UTF-8 (decode error at byte {exc.start}: {exc.reason}). "
            f"Re-encode (iconv / recode) or run a format-aware extractor first."
        ) from exc


def _extract_pdf(path: Path) -> tuple[str, int]:
    """Return (joined_text, page_count). Empty text on parser failure."""
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:
        raise UnsupportedFormat(
            "pypdf is required to read PDF sources. Install via the "
            "default project deps (`pip install stellars-claude-code-plugins`)."
        ) from exc
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise UnsupportedFormat(f"{path.name} cannot be parsed by pypdf: {exc}") from exc
    pages: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    parts: list[str] = []
    for i, p in enumerate(pages, start=1):
        if i > 1:
            parts.append(f"\n\n--- page {i} ---\n\n")
        parts.append(p)
    return "".join(parts), len(pages)


def _extract_docx(path: Path) -> str:
    """Iterate paragraphs + table cells via python-docx."""
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        raise UnsupportedFormat(
            "python-docx is required to read .docx sources. Install via "
            "default project deps (`pip install stellars-claude-code-plugins`)."
        ) from exc
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_odt(path: Path) -> str:
    """Walk ODT text:p / text:h elements via odfpy."""
    try:
        from odf import teletype  # type: ignore[import-not-found]
        from odf.opendocument import load  # type: ignore[import-not-found]
        from odf.text import H, P  # type: ignore[import-not-found]
    except ImportError as exc:
        raise UnsupportedFormat(
            "odfpy is required to read .odt sources. Install via default "
            "project deps (`pip install stellars-claude-code-plugins`)."
        ) from exc
    doc = load(str(path))
    parts: list[str] = []
    for element in doc.getElementsByType(H):
        parts.append(teletype.extractText(element))
    for element in doc.getElementsByType(P):
        parts.append(teletype.extractText(element))
    return "\n".join(p for p in parts if p)


def _extract_rtf(path: Path) -> str:
    """Strip RTF control words via striprtf."""
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore[import-not-found]
    except ImportError as exc:
        raise UnsupportedFormat(
            "striprtf is required to read .rtf sources. Install via default "
            "project deps (`pip install stellars-claude-code-plugins`)."
        ) from exc
    raw = path.read_text(encoding="utf-8", errors="replace")
    return rtf_to_text(raw)


class _HtmlTextExtractor(HTMLParser):
    """Strip-to-text using stdlib only (no bs4 dep)."""

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in {"script", "style"}:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style"} and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0:
            stripped = data.strip()
            if stripped:
                self.parts.append(stripped)


def _extract_html(path: Path) -> str:
    parser = _HtmlTextExtractor()
    parser.feed(path.read_text(encoding="utf-8", errors="replace"))
    return "\n".join(parser.parts)


# ---------------------------------------------------------------------------
# Language-inference helpers
# ---------------------------------------------------------------------------

# Tesseract uses 3-letter codes; langdetect returns ISO 639-1 (2-letter).
# Cover the common ones; anything else falls back to ``eng``.
_ISO2_TO_TESSERACT: dict[str, str] = {
    "en": "eng",
    "de": "deu",
    "fr": "fra",
    "es": "spa",
    "it": "ita",
    "pt": "por",
    "nl": "nld",
    "sv": "swe",
    "no": "nor",
    "da": "dan",
    "fi": "fin",
    "pl": "pol",
    "cs": "ces",
    "ru": "rus",
    "uk": "ukr",
    "ja": "jpn",
    "zh-cn": "chi_sim",
    "zh-tw": "chi_tra",
    "ko": "kor",
    "ar": "ara",
    "tr": "tur",
}


_FILENAME_LANG_TOKENS: dict[str, str] = {
    "-en": "eng",
    "-de": "deu",
    "-fr": "fra",
    "-es": "spa",
    "-it": "ita",
    "-pt": "por",
    "-pl": "pol",
    "-ru": "rus",
    "-jp": "jpn",
    "-cn": "chi_sim",
    "-kr": "kor",
    "_en": "eng",
    "_de": "deu",
    "_fr": "fra",
    "_es": "spa",
}


def _language_from_filename(name: str) -> str | None:
    """Pick up ``-de``, ``_fr`` style language hints in a filename."""
    for token, lang in _FILENAME_LANG_TOKENS.items():
        if token in name:
            return lang
    return None
