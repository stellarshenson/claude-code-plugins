"""Tests for stellars_claude_code_plugins.document_processing."""

from __future__ import annotations

import json

import pytest

from stellars_claude_code_plugins.document_processing import (
    GroundingMatch,
    ground,
    ground_many,
)
from stellars_claude_code_plugins.document_processing import settings as settings_mod
from stellars_claude_code_plugins.document_processing.chunking import (
    Chunk,
    recursive_chunk,
)
from stellars_claude_code_plugins.document_processing.cli import main as cli_main
from stellars_claude_code_plugins.document_processing.grounding import Location


class TestExactMatching:
    """Regex (exact) layer — whitespace-tolerant, case-insensitive."""

    def test_exact_verbatim(self):
        m = ground("quick brown fox", ["The quick brown fox jumps."])
        assert m.match_type == "exact"
        assert m.exact_score == 1.0
        assert m.exact_matched_text == "quick brown fox"
        assert m.exact_location.char_start == 4
        assert m.exact_location.char_end == 19

    def test_exact_case_insensitive(self):
        m = ground("QUICK BROWN FOX", ["The quick brown fox jumps."])
        assert m.match_type == "exact"
        assert m.exact_score == 1.0

    def test_exact_whitespace_tolerant(self):
        m = ground("quick brown fox", ["The  quick\n brown  \tfox jumps."])
        assert m.match_type == "exact"
        assert m.exact_score == 1.0

    def test_exact_miss(self):
        m = ground("completely unrelated phrase", ["The quick brown fox jumps."])
        assert m.exact_score == 0.0
        assert m.exact_matched_text == ""

    def test_exact_multi_source_first_hit_wins(self):
        m = ground(
            "brown fox", ["nothing here", "The quick brown fox jumps.", "also has brown fox"]
        )
        assert m.match_type == "exact"
        assert m.exact_location.source_index == 1

    def test_exact_with_source_paths(self):
        m = ground(
            "brown fox",
            [("doc1.txt", "nothing here"), ("doc2.txt", "The quick brown fox jumps.")],
        )
        assert m.match_type == "exact"
        assert m.exact_location.source_path == "doc2.txt"


class TestFuzzyMatching:
    """Levenshtein (fuzzy) layer — always runs, best across sources."""

    def test_fuzzy_above_threshold(self):
        m = ground(
            "quick brown fox jumped over",
            ["The quick brown fox jumps over the lazy dog."],
            fuzzy_threshold=0.80,
        )
        assert m.exact_score == 0.0
        assert m.fuzzy_score >= 0.80
        assert m.match_type == "fuzzy"

    def test_fuzzy_below_threshold(self):
        m = ground(
            "tropical island paradise",
            ["The quick brown fox jumps over the lazy dog."],
            fuzzy_threshold=0.85,
        )
        assert m.exact_score == 0.0
        assert m.fuzzy_score < 0.85
        assert m.match_type == "none"

    def test_fuzzy_always_computed_even_on_exact_hit(self):
        """Both scores always populated; exact match yields fuzzy=1.0 too."""
        m = ground("brown fox", ["The quick brown fox jumps."])
        assert m.exact_score == 1.0
        assert m.fuzzy_score == 1.0
        assert m.match_type == "exact"

    def test_fuzzy_best_across_sources(self):
        m = ground(
            "quick brown fox",
            [
                "blue sky overhead",
                "quirk brown fux jumps",
            ],
        )
        assert m.exact_score == 0.0
        assert m.fuzzy_score > 0.5
        assert m.fuzzy_location.source_index == 1


class TestBothSignalsReported:
    """All three scores always in the result (user requirement)."""

    def test_none_match_still_reports_fuzzy_signal(self):
        """Even when match_type=none, fuzzy_score shows best-effort signal."""
        m = ground("something different", ["slightly different content here"])
        assert m.match_type == "none"
        assert m.fuzzy_score > 0
        assert m.fuzzy_matched_text != ""

    def test_all_three_scores_independent(self):
        """exact=0 does not zero fuzzy or bm25."""
        m = ground("fox jumps", ["the quick fux jumps high"])
        assert m.exact_score == 0.0
        assert m.fuzzy_score > 0
        # BM25 may or may not fire on such a short source, but score is set

    def test_combined_score_is_max_of_three(self):
        m = ground("brown fox", ["The quick brown fox jumps."])
        assert m.combined_score == max(m.exact_score, m.fuzzy_score, m.bm25_score)


class TestBM25Matching:
    """BM25 layer — topical/lexical grounding across passages."""

    _LONG_SOURCE = (
        "Introduction paragraph about birds.\n\n"
        "The quick brown fox jumps over the lazy dog in the meadow.\n\n"
        "Cats sleep most of the day on windowsills.\n\n"
        "Aquatic mammals like dolphins are highly intelligent.\n"
    )

    def test_bm25_finds_right_passage(self):
        """Paraphrased claim with same key terms lands in the right passage."""
        m = ground(
            "fox and dog in a meadow",
            [self._LONG_SOURCE],
            fuzzy_threshold=0.95,  # high, so fuzzy fails
            bm25_threshold=0.4,
        )
        # The fox passage should win
        assert "fox" in m.bm25_matched_text
        assert m.bm25_score > 0
        assert m.bm25_token_recall > 0

    def test_bm25_token_recall_is_fraction(self):
        """Token recall = fraction of unique claim tokens in best passage."""
        m = ground("fox dog meadow", [self._LONG_SOURCE])
        # All 3 tokens present → recall = 1.0
        assert m.bm25_token_recall == 1.0

    def test_bm25_raw_score_available(self):
        """Raw BM25 score exposed for callers who want the unbounded signal."""
        m = ground("fox dog meadow", [self._LONG_SOURCE])
        assert m.bm25_raw_score >= 0

    def test_bm25_location_populated(self):
        """BM25 location has line/paragraph/page just like other layers."""
        m = ground("fox dog", [self._LONG_SOURCE])
        assert m.bm25_location.line_start > 0
        assert m.bm25_location.paragraph > 0
        assert m.bm25_location.page == 1

    def test_bm25_matches_topical_paraphrase(self):
        """Same terms, different order — BM25 catches what Levenshtein misses."""
        # "Dolphins are smart aquatic mammals" — paraphrase of sentence in source
        m = ground(
            "dolphins mammals intelligent aquatic",
            [self._LONG_SOURCE],
            fuzzy_threshold=0.95,
            bm25_threshold=0.5,
        )
        assert m.exact_score == 0.0
        # BM25 should catch this
        assert m.bm25_token_recall >= 0.5
        assert "dolphins" in m.bm25_matched_text.lower()

    def test_bm25_below_threshold_classified_none(self):
        """When BM25 token-recall below threshold, match_type=none."""
        m = ground(
            "quantum physics neutrino detector",
            [self._LONG_SOURCE],
            fuzzy_threshold=0.95,
            bm25_threshold=0.5,
        )
        assert m.match_type == "none"

    def test_bm25_priority_below_fuzzy(self):
        """When both fuzzy and bm25 would classify, fuzzy wins."""
        m = ground(
            "quick brown fox jumped",  # fuzzy match of "quick brown fox jumps"
            [self._LONG_SOURCE],
            fuzzy_threshold=0.80,
            bm25_threshold=0.5,
        )
        assert m.match_type == "fuzzy"  # fuzzy wins over bm25


class TestLocation:
    """Location metadata — line, column, paragraph, page, context."""

    def test_line_number_single_line(self):
        m = ground("fox", ["The quick brown fox jumps."])
        assert m.exact_location.line_start == 1
        assert m.exact_location.line_end == 1

    def test_line_number_multiline_source(self):
        text = "line one\nline two\nthe fox is here\nline four"
        m = ground("fox is here", [text])
        assert m.match_type == "exact"
        assert m.exact_location.line_start == 3
        assert m.exact_location.line_end == 3

    def test_column_number(self):
        text = "hello brown fox and more"
        m = ground("brown fox", [text])
        assert m.exact_location.line_start == 1
        # "brown fox" starts at char 6 on line 1 → column 7 (1-indexed)
        assert m.exact_location.column_start == 7

    def test_paragraph_number(self):
        text = "first paragraph text\n\nsecond paragraph with fox here\n\nthird paragraph"
        m = ground("fox", [text])
        assert m.match_type == "exact"
        assert m.exact_location.paragraph == 2

    def test_paragraph_blank_line_with_whitespace(self):
        """Blank lines with whitespace still separate paragraphs."""
        text = "first para\n  \t  \nsecond para fox"
        m = ground("fox", [text])
        assert m.exact_location.paragraph == 2

    def test_page_number_via_form_feed(self):
        """Pages separated by \\f (pdftotext convention)."""
        text = "page one content\fpage two with fox\fpage three"
        m = ground("fox", [text])
        assert m.match_type == "exact"
        assert m.exact_location.page == 2

    def test_page_1_when_no_form_feed(self):
        m = ground("fox", ["no form feeds here just fox content"])
        assert m.exact_location.page == 1

    def test_context_before_after(self):
        text = "The quick brown fox jumps over the lazy dog gently."
        m = ground("brown fox", [text])
        # Context should include surrounding words
        ctx_before = m.exact_location.context_before
        ctx_after = m.exact_location.context_after
        assert "quick" in ctx_before or "The" in ctx_before
        assert "jumps" in ctx_after or "over" in ctx_after

    def test_context_trimmed_to_max_chars(self):
        """Long context is trimmed with ellipsis."""
        long_text = "x" * 200 + " brown fox " + "y" * 200
        m = ground("brown fox", [long_text], context_chars=40)
        assert len(m.exact_location.context_before) <= 41  # 40 + ellipsis
        assert len(m.exact_location.context_after) <= 41


class TestEdgeCases:
    def test_empty_sources(self):
        m = ground("anything", [])
        assert isinstance(m, GroundingMatch)
        assert m.match_type == "none"
        assert m.exact_score == 0.0
        assert m.fuzzy_score == 0.0

    def test_empty_claim(self):
        m = ground("", ["some source text"])
        assert m.exact_score == 0.0

    def test_empty_source_text(self):
        m = ground("anything", [""])
        assert m.exact_score == 0.0
        assert m.fuzzy_score == 0.0

    def test_location_dataclass_default_is_neg_one(self):
        loc = Location()
        assert loc.line_start == -1
        assert loc.paragraph == -1
        assert loc.page == -1


class TestBatch:
    def test_ground_many_preserves_order(self):
        claims = ["brown fox", "lazy dog", "unrelated claim"]
        sources = ["The quick brown fox jumps over the lazy dog."]
        results = ground_many(claims, sources, fuzzy_threshold=0.85)
        assert len(results) == 3
        assert results[0].match_type == "exact"
        assert results[1].match_type == "exact"
        assert results[2].match_type in ("fuzzy", "none")


class TestCLI:
    """End-to-end CLI tests via the main() entrypoint."""

    def test_caveman_catalogue(self, capsys):
        """`document-processing --caveman` prints every subcommand with a terse one-liner."""
        code = cli_main(["--caveman"])
        assert code == 0
        out = capsys.readouterr().out
        assert "(caveman)" in out
        for sub, expected in (
            ("ground", "one claim"),
            ("batch-ground", "many claims"),
            ("extract-claims", "doc -> claims.json"),
            ("check-consistency", "doc fight self"),
            ("batch-validate", "yaml manifest"),
            ("setup", "first run"),
        ):
            assert sub in out, f"{sub} missing from caveman catalogue"
            assert expected in out, f"caveman line for {sub} missing"

    def test_ground_exit_zero_on_hit(self, tmp_path, capsys):
        src = tmp_path / "src.txt"
        src.write_text("The quick brown fox jumps over the lazy dog.")
        code = cli_main(["ground", "--claim", "brown fox", "--source", str(src)])
        assert code == 0
        out = capsys.readouterr().out
        assert "EXACT" in out
        assert "exact=1.000" in out

    def test_ground_output_includes_location(self, tmp_path, capsys):
        src = tmp_path / "src.txt"
        src.write_text("line one\nline two with brown fox here\nline three")
        code = cli_main(["ground", "--claim", "brown fox", "--source", str(src)])
        assert code == 0
        out = capsys.readouterr().out
        assert "L2" in out  # line 2
        assert "¶1" in out  # paragraph 1

    def test_ground_exit_one_on_miss(self, tmp_path, capsys):
        src = tmp_path / "src.txt"
        src.write_text("Only about horticulture.")
        code = cli_main(
            [
                "ground",
                "--claim",
                "quantum physics",
                "--source",
                str(src),
                "--threshold",
                "0.95",
                "--semantic",
                "off",  # ensure test doesn't depend on optional model download
            ]
        )
        assert code == 1
        out = capsys.readouterr().out
        assert "NONE" in out

    def test_ground_json_output(self, tmp_path, capsys):
        src = tmp_path / "src.txt"
        src.write_text("The quick brown fox jumps.")
        code = cli_main(["ground", "--claim", "brown fox", "--source", str(src), "--json"])
        assert code == 0
        out = capsys.readouterr().out
        data = json.loads(out)
        assert data["match_type"] == "exact"
        assert data["exact_score"] == 1.0
        assert data["fuzzy_score"] >= 0.0
        # Nested Location should serialize too
        assert "exact_location" in data
        assert data["exact_location"]["line_start"] == 1
        assert data["exact_location"]["paragraph"] == 1

    def test_ground_many_markdown_report(self, tmp_path):
        src = tmp_path / "src.txt"
        src.write_text("The quick brown fox jumps over the lazy dog.")
        claims = tmp_path / "claims.json"
        claims.write_text(json.dumps(["brown fox", "unrelated claim"]))
        out = tmp_path / "report.md"
        cli_main(
            [
                "batch-ground",
                "--claims",
                str(claims),
                "--source",
                str(src),
                "--output",
                str(out),
            ]
        )
        report = out.read_text()
        assert "Grounding Report" in report
        assert "CONFIRMED" in report
        assert "L1" in report  # location in report
        assert "¶1" in report

    def test_ground_many_json_report(self, tmp_path):
        src = tmp_path / "src.txt"
        src.write_text("The quick brown fox jumps.")
        claims = tmp_path / "claims.json"
        claims.write_text(json.dumps([{"claim": "brown fox"}, {"claim": "missing"}]))
        out = tmp_path / "report.json"
        cli_main(
            [
                "batch-ground",
                "--claims",
                str(claims),
                "--source",
                str(src),
                "--output",
                str(out),
                "--json",
            ]
        )
        data = json.loads(out.read_text())
        assert data["summary"]["total"] == 2
        assert data["summary"]["exact"] >= 1
        assert len(data["matches"]) == 2
        # Location fields should be in the JSON
        assert "exact_location" in data["matches"][0]

    def test_ground_missing_source_errors(self, capsys):
        with pytest.raises(SystemExit):
            cli_main(["ground", "--claim", "anything", "--source", "/nonexistent/file.txt"])


class TestChunking:
    """Recursive chunking preserves offsets + boundaries."""

    def test_empty_text_returns_empty(self):
        assert recursive_chunk("") == []

    def test_short_text_one_chunk(self):
        text = "The quick brown fox."
        chunks = recursive_chunk(text, max_chars=1500)
        assert len(chunks) == 1
        assert chunks[0].text == text
        assert chunks[0].char_start == 0
        assert chunks[0].char_end == len(text)

    def test_paragraph_split(self):
        text = "First paragraph here.\n\nSecond paragraph longer content here.\n\nThird paragraph ends the document."
        chunks = recursive_chunk(text, max_chars=30, min_chunk_chars=10)
        assert len(chunks) >= 2

    def test_offsets_are_valid(self):
        text = "paragraph one\n\nparagraph two content\n\nparagraph three final"
        chunks = recursive_chunk(text, max_chars=200, min_chunk_chars=5)
        for c in chunks:
            # Char offsets must be valid bounds into the source
            assert 0 <= c.char_start < c.char_end <= len(text)
            # First + last words of the chunk should appear inside the source span
            first_word = c.text.split()[0]
            last_word = c.text.split()[-1]
            span = text[c.char_start : c.char_end]
            assert first_word in span
            assert last_word in span

    def test_long_sentence_sliding_window(self):
        sentence = "word " * 500  # long single sentence
        chunks = recursive_chunk(sentence, max_chars=200, overlap_chars=50)
        assert len(chunks) > 1
        # Overlap: consecutive chunks share some content
        if len(chunks) >= 2:
            assert chunks[0].char_end > chunks[1].char_start

    def test_offsets_monotonic(self):
        text = "a " * 200 + "\n\n" + "b " * 200
        chunks = recursive_chunk(text, max_chars=100)
        starts = [c.char_start for c in chunks]
        assert starts == sorted(starts)

    def test_chunk_dataclass(self):
        c = Chunk("hello", 0, 5)
        assert len(c) == 5


class TestSettings:
    """Settings load/save/prompt — zero-dep."""

    def test_defaults_when_no_file(self, tmp_path, monkeypatch):
        monkeypatch.chdir(tmp_path)
        monkeypatch.setenv("HOME", str(tmp_path))
        cfg = settings_mod.load()
        assert cfg.semantic_enabled is False
        assert cfg.semantic_model == "intfloat/multilingual-e5-small"
        assert cfg.semantic_device == "auto"

    def test_save_then_load_roundtrip(self, tmp_path, monkeypatch):
        monkeypatch.chdir(tmp_path)
        monkeypatch.setenv("HOME", str(tmp_path))
        # Create project root marker
        (tmp_path / ".claude").mkdir()
        s = settings_mod.Settings(semantic_enabled=True, semantic_model="custom/model")
        path = settings_mod.save(s)
        assert path.exists()
        loaded = settings_mod.load()
        assert loaded.semantic_enabled is True
        assert loaded.semantic_model == "custom/model"

    def test_is_semantic_available_reflects_imports(self):
        # This may be True or False depending on test env — just ensure the
        # function runs without error and returns a bool
        result = settings_mod.is_semantic_available()
        assert isinstance(result, bool)

    def test_install_hint_is_helpful(self):
        hint = settings_mod.semantic_install_hint()
        assert "pip install" in hint
        assert "semantic" in hint.lower()


class TestSemanticOnContract:
    """`--semantic on` is an explicit contract:

    - deps present -> grounder is built, returned
    - deps missing -> hard fail (sys.exit 2), do NOT silently degrade

    `--semantic` not passed (config default takes over) keeps the old
    warn-and-degrade behaviour for back-compat. `--semantic off` is
    handled earlier and never reaches the deps-check.
    """

    def test_semantic_on_with_deps_missing_exits_2(self, monkeypatch, capsys):
        """The bug we saw in the 49th-batch grounding: --semantic on,
        deps missing, CLI proceeded silently. New contract: hard fail.
        """
        from stellars_claude_code_plugins.document_processing.cli import (
            _build_semantic_grounder,
        )

        monkeypatch.setattr(settings_mod, "is_semantic_available", lambda: False)
        cfg = settings_mod.Settings(semantic_enabled=False)

        with pytest.raises(SystemExit) as exc_info:
            _build_semantic_grounder(cfg, cli_override="on")
        assert exc_info.value.code == 2

        err = capsys.readouterr().err
        assert "ERROR: --semantic on requires the [semantic] extras" in err
        assert "pip install" in err

    def test_semantic_config_default_with_deps_missing_warns_and_degrades(
        self, monkeypatch, capsys
    ):
        """Config says enabled, deps missing, no explicit --semantic flag.
        Back-compat: warn, return None, let the caller continue without
        semantic. Don't surprise users with exit 2 just because their
        config has semantic_enabled=true on a machine without the extras.
        """
        from stellars_claude_code_plugins.document_processing.cli import (
            _build_semantic_grounder,
        )

        monkeypatch.setattr(settings_mod, "is_semantic_available", lambda: False)
        cfg = settings_mod.Settings(semantic_enabled=True)

        result = _build_semantic_grounder(cfg, cli_override=None)
        assert result is None

        err = capsys.readouterr().err
        assert "WARNING: semantic grounding enabled in settings" in err
        # NOT the explicit-fail error
        assert "ERROR: --semantic on" not in err

    def test_semantic_off_short_circuits_before_deps_check(self, monkeypatch):
        """--semantic off returns None before the deps check; setting the
        availability to False should not matter because we never get there.
        """
        from stellars_claude_code_plugins.document_processing.cli import (
            _build_semantic_grounder,
        )

        sentinel_called = []
        monkeypatch.setattr(
            settings_mod,
            "is_semantic_available",
            lambda: sentinel_called.append(True) or False,
        )
        cfg = settings_mod.Settings(semantic_enabled=True)

        result = _build_semantic_grounder(cfg, cli_override="off")
        assert result is None
        assert sentinel_called == []  # deps-check never invoked

    def test_validate_many_helper_honours_same_contract(self, monkeypatch, capsys):
        """validate_many.py has a copy of the same function (to dodge a
        circular import) - it must apply the same hard-fail rule.
        """
        from stellars_claude_code_plugins.document_processing.validate_many import (
            _maybe_build_grounder,
        )

        monkeypatch.setattr(settings_mod, "is_semantic_available", lambda: False)
        cfg = settings_mod.Settings(semantic_enabled=False)

        with pytest.raises(SystemExit) as exc_info:
            _maybe_build_grounder(cfg, semantic_override="on")
        assert exc_info.value.code == 2

        err = capsys.readouterr().err
        assert "ERROR: --semantic on requires the [semantic] extras" in err


class TestCLISetup:
    """CLI setup subcommand."""

    def test_setup_shows_current_if_present(self, tmp_path, monkeypatch, capsys):
        monkeypatch.chdir(tmp_path)
        monkeypatch.setenv("HOME", str(tmp_path))
        (tmp_path / ".claude").mkdir()
        # Pre-seed settings
        settings_mod.save(settings_mod.Settings(semantic_enabled=True))
        code = cli_main(["setup"])
        assert code == 0
        err = capsys.readouterr().err
        assert "semantic_enabled" in err


# -------------------------------------------------------------------------
# Source-format fallback: gate-warning on unsupported binaries (was WI#1
# binary-rejection; Release F changed the contract to skip-with-warning).
# -------------------------------------------------------------------------


class TestBinarySourceFallback:
    """Binary inputs no longer hard-block with exit 2 - they emit a SKIPPED
    warning through the stop-and-think gate. The gate exits 2 by default
    (warning unacked); the agent acks per source with `--ack-warning
    TOKEN='skip-binary-source-acceptable'` to proceed.
    """

    def test_png_magic_bytes_emit_skipped_gate_warning(self, tmp_path, capsys):
        png_bytes = b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
        p = tmp_path / "weird.txt"  # renamed PNG
        p.write_bytes(png_bytes)
        claim = tmp_path / "claim.json"
        claim.write_text(json.dumps(["hello world"]))
        with pytest.raises(SystemExit) as exc:
            cli_main(["batch-ground", "--claims", str(claim), "--source", str(p)])
        # Gate exits 2 with the unacked SOURCE-SKIPPED warning.
        assert exc.value.code == 2
        err = capsys.readouterr().err
        assert "BLOCKED" in err
        assert "SOURCE-SKIPPED" in err

    def test_pdf_with_invalid_bytes_skipped(self, tmp_path, capsys):
        # A PDF that pypdf cannot parse triggers UnsupportedFormat which
        # routes to SOURCE-SKIPPED through the gate.
        p = tmp_path / "document.pdf"
        p.write_bytes(b"%PDF-1.4 fake content here")
        with pytest.raises(SystemExit) as exc:
            cli_main(["ground", "--claim", "anything", "--source", str(p)])
        assert exc.value.code == 2
        err = capsys.readouterr().err
        assert "BLOCKED" in err

    def test_invalid_utf8_skipped(self, tmp_path, capsys):
        p = tmp_path / "bad.txt"
        p.write_bytes(b"valid ascii then invalid: \xff\xfe\xfd")
        with pytest.raises(SystemExit) as exc:
            cli_main(["ground", "--claim", "anything", "--source", str(p)])
        assert exc.value.code == 2
        err = capsys.readouterr().err
        assert "BLOCKED" in err
        assert "SOURCE-SKIPPED" in err

    def test_plain_text_accepted(self, tmp_path):
        p = tmp_path / "plain.txt"
        p.write_text("hello world")
        code = cli_main(["ground", "--claim", "hello world", "--source", str(p)])
        assert code == 0


# -------------------------------------------------------------------------
# Release F: source-format fallback + sibling lookup + auto-OCR
# -------------------------------------------------------------------------


def _ack_all_warnings(stderr_out: str) -> list[str]:
    """Helper - extract every W-xxxxxxxx token from BLOCKED stderr and
    return a list of `--ack-warning TOKEN=test-fixture` argv pairs."""
    import re

    seen: set[str] = set()
    flags: list[str] = []
    for tok in re.findall(r"W-[0-9a-f]{8}", stderr_out):
        if tok in seen:
            continue
        seen.add(tok)
        flags += ["--ack-warning", f"{tok}=test-fixture"]
    return flags


class TestSourceFormats:
    """Source loading via extractors module: PDF text/scanned, DOCX,
    sibling lookup, OCR-LANG-NEEDED gate, OCR-MISSING gate, auto-OCR
    quality bands. Uses mocks for OCR + scanned-PDF detection so tests
    do not need a real Tesseract install."""

    def test_text_source_passthrough(self, tmp_path):
        """Plain .txt remains the simple happy-path."""
        p = tmp_path / "plain.txt"
        p.write_text("the brown fox jumps over the lazy dog")
        code = cli_main(["ground", "--claim", "the brown fox", "--source", str(p)])
        assert code == 0

    def test_docx_source_extracted(self, tmp_path):
        """python-docx produces a .docx; grounding finds the claim."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("The brown fox jumps over the lazy dog.")
        doc.add_paragraph("Second paragraph with more content.")
        p = tmp_path / "report.docx"
        doc.save(str(p))
        code = cli_main(["ground", "--claim", "the brown fox", "--source", str(p)])
        assert code == 0

    def test_pdf_text_extraction(self, tmp_path, monkeypatch):
        """Mock extract_text to return pdf-text kind; grounding succeeds."""
        from stellars_claude_code_plugins.document_processing import extractors

        p = tmp_path / "document.pdf"
        p.write_bytes(b"%PDF-1.4 stub")

        def fake_extract(path, scanned_threshold=100):
            return extractors.Extracted(
                text="the brown fox jumps over the lazy dog",
                kind="pdf-text",
                page_count=1,
            )

        monkeypatch.setattr(extractors, "extract_text", fake_extract)
        code = cli_main(["ground", "--claim", "the brown fox", "--source", str(p)])
        assert code == 0

    def test_scanned_pdf_with_ocr_txt_sibling(self, tmp_path, monkeypatch, capsys):
        """Scanned PDF + sibling .ocr.txt with header → OCR-CANDIDATE warning."""
        from stellars_claude_code_plugins.document_processing import extractors

        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 stub")
        sibling = tmp_path / "scanned.ocr.txt"
        sibling.write_text(
            "# OCR candidate for scanned.pdf\n"
            "# quality: candidate (mean conf 67%, 2 pages, 100 chars)\n"
            "# lang: eng\n"
            "# generated: 2026-04-27 10:00 UTC\n"
            "# NOTE: review etc\n"
            "\n"
            "the brown fox jumps over the lazy dog"
        )

        original = extractors.extract_text

        def fake_extract(path, scanned_threshold=100):
            if path.suffix.lower() == ".pdf":
                return extractors.Extracted(text="", kind="pdf-scanned", page_count=2)
            return original(path, scanned_threshold=scanned_threshold)

        monkeypatch.setattr(extractors, "extract_text", fake_extract)

        # First run blocks (gate fires OCR-CANDIDATE for unreviewed sibling).
        with pytest.raises(SystemExit) as exc:
            cli_main(["ground", "--claim", "the brown fox", "--source", str(pdf)])
        assert exc.value.code == 2
        err = capsys.readouterr().err
        assert "OCR-CANDIDATE" in err
        assert "scanned.ocr.txt" in err

        # Ack with reason → grounding succeeds via sibling text.
        ack_flags = _ack_all_warnings(err)
        code = cli_main(["ground", "--claim", "the brown fox", "--source", str(pdf), *ack_flags])
        assert code == 0

    def test_scanned_pdf_with_reviewed_ocr_sibling_no_candidate_warn(self, tmp_path, monkeypatch):
        """Sibling without header (human-reviewed) fires OCR-FALLBACK not
        OCR-CANDIDATE - candidate has graduated to ground truth."""
        from stellars_claude_code_plugins.document_processing import extractors

        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 stub")
        sibling = tmp_path / "scanned.ocr.txt"
        sibling.write_text("the brown fox jumps over the lazy dog")  # no header

        original = extractors.extract_text

        def fake_extract(path, scanned_threshold=100):
            if path.suffix.lower() == ".pdf":
                return extractors.Extracted(text="", kind="pdf-scanned", page_count=2)
            return original(path, scanned_threshold=scanned_threshold)

        monkeypatch.setattr(extractors, "extract_text", fake_extract)
        # Single ack call - the sibling fires OCR-FALLBACK rather than CANDIDATE.
        with pytest.raises(SystemExit) as exc:
            cli_main(["ground", "--claim", "the brown fox", "--source", str(pdf)])
        # First run blocks with OCR-FALLBACK
        assert exc.value.code == 2

    def test_scanned_pdf_no_ocr_lang_emits_lang_needed(self, tmp_path, monkeypatch, capsys):
        """Scanned PDF + no sibling + no --ocr-lang → OCR-LANG-NEEDED."""
        from stellars_claude_code_plugins.document_processing import extractors

        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 stub")

        def fake_extract(path, scanned_threshold=100):
            return extractors.Extracted(text="", kind="pdf-scanned", page_count=2)

        monkeypatch.setattr(extractors, "extract_text", fake_extract)

        with pytest.raises(SystemExit) as exc:
            cli_main(["ground", "--claim", "anything", "--source", str(pdf)])
        assert exc.value.code == 2
        err = capsys.readouterr().err
        assert "OCR-LANG-NEEDED" in err
        assert "--ocr-lang" in err

    def test_scanned_pdf_with_ocr_lang_no_deps_emits_missing(self, tmp_path, monkeypatch, capsys):
        """Scanned PDF + --ocr-lang + no OCR deps → OCR-MISSING."""
        from stellars_claude_code_plugins.document_processing import (
            extractors,
        )
        from stellars_claude_code_plugins.document_processing import (
            ocr as ocr_mod,
        )

        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 stub")

        def fake_extract(path, scanned_threshold=100):
            return extractors.Extracted(text="", kind="pdf-scanned", page_count=2)

        monkeypatch.setattr(extractors, "extract_text", fake_extract)
        monkeypatch.setattr(ocr_mod, "ocr_available", lambda: False)

        with pytest.raises(SystemExit) as exc:
            cli_main(["ground", "--claim", "anything", "--source", str(pdf), "--ocr-lang", "eng"])
        assert exc.value.code == 2
        err = capsys.readouterr().err
        assert "OCR-MISSING" in err

    def test_auto_ocr_good_quality_caches_and_passes(self, tmp_path, monkeypatch, capsys):
        """Mock ocr_pdf to return quality=good; cache file written; source flows."""
        from stellars_claude_code_plugins.document_processing import (
            extractors,
        )
        from stellars_claude_code_plugins.document_processing import (
            ocr as ocr_mod,
        )

        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 stub")

        def fake_extract(path, scanned_threshold=100):
            if path.suffix == ".pdf":
                return extractors.Extracted(text="", kind="pdf-scanned", page_count=2)
            return extractors.Extracted(
                text=path.read_text(encoding="utf-8", errors="strict"), kind="text"
            )

        monkeypatch.setattr(extractors, "extract_text", fake_extract)
        monkeypatch.setattr(ocr_mod, "ocr_available", lambda: True)

        good_text = "the brown fox jumps over the lazy dog " * 5

        def fake_ocr_pdf(path, lang, *, dpi=200):
            return ocr_mod.OcrResult(
                text=good_text,
                mean_confidence=92.0,
                per_page_confidence=[91.0, 93.0],
                total_chars=len(good_text),
                quality="good",
                lang=lang,
            )

        monkeypatch.setattr(ocr_mod, "ocr_pdf", fake_ocr_pdf)

        # First run blocks with OCR-FALLBACK; cache file is written
        # alongside even though the gate blocked the grounding step.
        with pytest.raises(SystemExit) as exc:
            cli_main(
                [
                    "ground",
                    "--claim",
                    "the brown fox",
                    "--source",
                    str(pdf),
                    "--ocr-lang",
                    "eng",
                ]
            )
        assert exc.value.code == 2
        err = capsys.readouterr().err
        assert "OCR-FALLBACK" in err
        assert "scanned.ocr.txt" in err

        # Cache file written with the OCR header.
        cache = tmp_path / "scanned.ocr.txt"
        assert cache.exists()
        assert "# OCR candidate for scanned.pdf" in cache.read_text()
        assert "the brown fox" in cache.read_text()

        # Simulate agent reviewing the candidate by stripping the header.
        # On rerun the gate should fire OCR-FALLBACK (sibling, not
        # candidate) and the source flows through to grounding.
        cache.write_text(good_text)

        # Capture the (different) OCR-FALLBACK token from a fresh run
        # without acks, then rerun with the matching ack.
        with pytest.raises(SystemExit):
            cli_main(
                [
                    "ground",
                    "--claim",
                    "the brown fox",
                    "--source",
                    str(pdf),
                    "--ocr-lang",
                    "eng",
                ]
            )
        err2 = capsys.readouterr().err
        ack_flags = _ack_all_warnings(err2)
        code = cli_main(
            [
                "ground",
                "--claim",
                "the brown fox",
                "--source",
                str(pdf),
                "--ocr-lang",
                "eng",
                *ack_flags,
            ]
        )
        assert code == 0

    def test_auto_ocr_failed_quality_skips_source(self, tmp_path, monkeypatch, capsys):
        """Mock OCR returns quality=failed; OCR-FAILED warning + source skipped."""
        from stellars_claude_code_plugins.document_processing import (
            extractors,
        )
        from stellars_claude_code_plugins.document_processing import (
            ocr as ocr_mod,
        )

        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 stub")

        def fake_extract(path, scanned_threshold=100):
            return extractors.Extracted(text="", kind="pdf-scanned", page_count=2)

        monkeypatch.setattr(extractors, "extract_text", fake_extract)
        monkeypatch.setattr(ocr_mod, "ocr_available", lambda: True)

        def fake_ocr_pdf(path, lang, *, dpi=200):
            return ocr_mod.OcrResult(
                text="garbled",
                mean_confidence=42.0,
                per_page_confidence=[42.0],
                total_chars=7,
                quality="failed",
                failure_reason="mean confidence 42.0% < 60%",
                lang=lang,
            )

        monkeypatch.setattr(ocr_mod, "ocr_pdf", fake_ocr_pdf)

        with pytest.raises(SystemExit) as exc:
            cli_main(
                [
                    "ground",
                    "--claim",
                    "anything",
                    "--source",
                    str(pdf),
                    "--ocr-lang",
                    "eng",
                ]
            )
        assert exc.value.code == 2
        err = capsys.readouterr().err
        assert "OCR-FAILED" in err
        assert "vision-OCR" in err.lower() or "Read tool" in err
        # Candidate cached for editing even though source was skipped.
        assert (tmp_path / "scanned.ocr.txt").exists()

    def test_sibling_priority_ocr_txt_beats_txt_beats_docx(self, tmp_path):
        """find_sibling returns highest-priority match."""
        from stellars_claude_code_plugins.document_processing.extractors import find_sibling

        pdf = tmp_path / "report.pdf"
        pdf.write_bytes(b"%PDF-1.4")
        (tmp_path / "report.docx").write_bytes(b"PK\x03\x04 stub")
        (tmp_path / "report.txt").write_text("text")
        (tmp_path / "report.ocr.txt").write_text("ocr text")

        assert find_sibling(pdf).name == "report.ocr.txt"

    def test_sibling_excludes_image_extensions(self, tmp_path):
        """A sibling .png next to a .pdf is NOT a candidate sibling."""
        from stellars_claude_code_plugins.document_processing.extractors import find_sibling

        pdf = tmp_path / "report.pdf"
        pdf.write_bytes(b"%PDF-1.4")
        (tmp_path / "report.png").write_bytes(b"\x89PNG\r\n\x1a\n")

        assert find_sibling(pdf) is None

    def test_language_inference_from_filename(self, tmp_path):
        """Filename hints (-de, _fr) feed the language suggestion."""
        from stellars_claude_code_plugins.document_processing.extractors import (
            infer_pdf_language,
        )

        de_path = tmp_path / "rapport-de.pdf"
        de_path.write_bytes(b"%PDF-1.4")
        assert infer_pdf_language("", de_path) == "deu"

        fr_path = tmp_path / "rapport_fr.pdf"
        fr_path.write_bytes(b"%PDF-1.4")
        assert infer_pdf_language("", fr_path) == "fra"

        en_path = tmp_path / "report.pdf"
        en_path.write_bytes(b"%PDF-1.4")
        assert infer_pdf_language("", en_path) == "eng"  # default

    def test_ocr_candidate_header_format(self, tmp_path):
        """cache_ocr_candidate writes a header with quality / lang / timestamp."""
        from stellars_claude_code_plugins.document_processing import ocr as ocr_mod

        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4")

        result = ocr_mod.OcrResult(
            text="the brown fox",
            mean_confidence=72.0,
            per_page_confidence=[72.0],
            total_chars=13,
            quality="candidate",
            lang="eng",
        )
        cache = ocr_mod.cache_ocr_candidate(pdf, result)
        body = cache.read_text()

        assert "# OCR candidate for scanned.pdf" in body
        assert "# quality: candidate" in body
        assert "# lang: eng" in body
        assert "# generated:" in body
        assert "the brown fox" in body
        # Header still in place → file is unreviewed.
        assert ocr_mod.has_unreviewed_header(cache)

        # Removing the header marks reviewed.
        cleaned_body = "the brown fox"
        cache.write_text(cleaned_body)
        assert not ocr_mod.has_unreviewed_header(cache)


# -------------------------------------------------------------------------
# WI#3: cross-source provenance + --primary-source
# -------------------------------------------------------------------------


class TestCrossSourceProvenance:
    def test_grounded_source_from_exact_hit(self):
        sources = [
            ("primary.md", "The cat sat on the mat."),
            ("secondary.md", "Unrelated content about dogs."),
        ]
        m = ground("the cat sat on the mat", sources)
        assert m.match_type == "exact"
        assert m.grounded_source == "primary.md"
        assert m.is_primary_source is True

    def test_non_primary_flag(self):
        sources = [
            ("primary.md", "Nothing about cats here."),
            ("secondary.md", "The cat sat on the mat."),
        ]
        m = ground("the cat sat on the mat", sources, primary_source="primary.md")
        assert m.grounded_source == "secondary.md"
        assert m.is_primary_source is False
        assert m.verification_needed is True

    def test_primary_source_match(self):
        sources = [
            ("primary.md", "The cat sat on the mat."),
            ("secondary.md", "Also: the cat sat on the mat."),
        ]
        m = ground("the cat sat on the mat", sources, primary_source="primary.md")
        assert m.is_primary_source is True


# -------------------------------------------------------------------------
# WI#5 + WI#6: lexical_co_support, verification_needed, claim_attributes
# -------------------------------------------------------------------------


class TestVerificationSignals:
    def test_exact_hit_has_lexical_support(self):
        m = ground("the cat sat", [("a.txt", "the cat sat on the mat")])
        assert m.lexical_co_support is True

    def test_claim_attributes_populated(self):
        m = ground(
            "42 users logged in yesterday",
            [("a.txt", "42 users logged in yesterday.")],
        )
        attrs = m.claim_attributes
        assert "numbers" in attrs
        assert "entities" in attrs
        assert "passage_numbers" in attrs
        assert "passage_entities" in attrs
        # At least the number 42 should be extracted
        values = [v for v, _, _ in attrs["numbers"]]
        assert "42" in values

    def test_numeric_co_presence_triggers_verification(self):
        # Both sides have numbers tied to the same context word ("users")
        # but the deterministic mismatch detector won't fire with a single
        # clean hit — the heuristic flag should still call this out.
        m = ground(
            "the project grew to 42 users",
            [("a.txt", "the project reports 100 users on record")],
        )
        # bm25 / fuzzy co-occurrence should yield verification_needed=True
        # because both claim and passage have number+"users"
        if m.match_type in ("fuzzy", "bm25", "semantic"):
            assert m.verification_needed is True


# -------------------------------------------------------------------------
# WI#2: extract-claims
# -------------------------------------------------------------------------


class TestExtractClaims:
    def test_basic_extraction(self, tmp_path):
        from stellars_claude_code_plugins.document_processing.extract import (
            extract_claims,
        )

        doc = (
            "# Heading\n\n"
            "The system handles 42 concurrent sessions. "
            "It was tested on Linux and macOS.\n\n"
            "Short.\n\n"
            "- dev\n"
            "- test\n"
            "- staging\n\n"
            "The deployment runs on Kubernetes with three nodes.\n"
        )
        claims = extract_claims(doc)
        assert len(claims) >= 2
        # Stable IDs
        assert claims[0].id.startswith("c0")
        # Short stubs and pure headers excluded
        for c in claims:
            assert len(c.claim) >= 20

    def test_cli_extract_claims(self, tmp_path, capsys):
        doc = tmp_path / "doc.md"
        doc.write_text(
            "The system handles 42 concurrent sessions.\nIt was tested on Linux and macOS.\n"
        )
        out = tmp_path / "claims.json"
        code = cli_main(["extract-claims", "--document", str(doc), "--output", str(out)])
        assert code == 0
        data = json.loads(out.read_text())
        assert len(data) >= 1
        assert "id" in data[0]
        assert "claim" in data[0]
        assert "line_number" in data[0]


# -------------------------------------------------------------------------
# WI#4: check-consistency
# -------------------------------------------------------------------------


class TestCheckConsistency:
    def test_numeric_divergence_flagged(self):
        from stellars_claude_code_plugins.document_processing.consistency import (
            check_consistency,
        )

        text = (
            "The platform supports 42 users on average.\n"
            "\n\n"
            "Recent benchmarks show 50 users on load.\n"
        )
        findings = check_consistency(text)
        numeric_findings = [f for f in findings if f.kind == "numeric"]
        assert len(numeric_findings) >= 1
        # Both line numbers should appear
        all_lines = [line for f in numeric_findings for line, _ in f.occurrences]
        assert 1 in all_lines
        assert 4 in all_lines

    def test_entity_set_divergence_flagged(self):
        from stellars_claude_code_plugins.document_processing.consistency import (
            check_consistency,
        )

        text = (
            "We run dev, test, and staging environments.\n"
            "\n\n"
            "Pipeline deploys to dev, staging, and prod.\n"
        )
        findings = check_consistency(text)
        set_findings = [f for f in findings if f.kind == "entity_set"]
        assert len(set_findings) >= 1

    def test_no_divergence_reports_clean(self):
        from stellars_claude_code_plugins.document_processing.consistency import (
            check_consistency,
            format_consistency_report,
        )

        text = "Simple consistent document with 42 users and 42 users again.\n"
        findings = check_consistency(text)
        # Same value twice - no divergence
        numeric_findings = [f for f in findings if f.kind == "numeric"]
        assert not numeric_findings
        report = format_consistency_report(findings)
        assert "No divergences" in report

    def test_cli_check_consistency(self, tmp_path):
        doc = tmp_path / "doc.md"
        doc.write_text(
            "The system handles 42 users per session.\n\nRecent tests show 50 users per session.\n"
        )
        out = tmp_path / "consistency.md"
        code = cli_main(["check-consistency", "--document", str(doc), "--output", str(out)])
        # Exit 1 when findings exist
        assert code == 1
        assert out.exists()
        report = out.read_text()
        assert "Self-Consistency" in report


# -------------------------------------------------------------------------
# WI#7: batch-validate
# -------------------------------------------------------------------------


class TestValidateMany:
    def test_basic_batch(self, tmp_path, monkeypatch):
        # Redirect HOME so settings don't leak from the real user
        monkeypatch.setenv("HOME", str(tmp_path / "fake_home"))
        (tmp_path / "fake_home").mkdir()

        # Build two tiny client fixtures
        client_a = tmp_path / "clients" / "alpha"
        client_a.mkdir(parents=True)
        (client_a / "transcript.md").write_text(
            "Alpha team is building an API for payment processing.\n"
            "They plan to launch in Q3 2026 with 5 team members.\n"
        )
        (client_a / "brief.md").write_text(
            "Alpha team builds a payment API.\n"
            "Launch planned for Q3 2026.\n"
            "Team size: 5 members.\n"
        )

        client_b = tmp_path / "clients" / "beta"
        client_b.mkdir(parents=True)
        (client_b / "transcript.md").write_text(
            "Beta team works on a data pipeline using Apache Spark.\n"
            "They handle 1M rows per day.\n"
        )
        (client_b / "brief.md").write_text(
            "Beta team runs a data pipeline on Apache Spark.\nDaily volume is 1M rows.\n"
        )

        source_map = tmp_path / "source_map.yaml"
        source_map.write_text(
            "clients:\n"
            "  alpha:\n"
            "    sources: [clients/alpha/transcript.md]\n"
            "    document: clients/alpha/brief.md\n"
            "  beta:\n"
            "    sources: [clients/beta/transcript.md]\n"
            "    document: clients/beta/brief.md\n"
        )

        output_dir = tmp_path / "validation"
        code = cli_main(
            [
                "batch-validate",
                "--source-map",
                str(source_map),
                "--output-dir",
                str(output_dir),
                "--semantic",
                "off",
            ]
        )
        # Code is 0 only if everything grounded AND no inconsistencies.
        # On a tiny fixture some claims may not reach thresholds; accept 0 or 1.
        assert code in (0, 1)
        assert (output_dir / "alpha" / "grounding-report.md").exists()
        assert (output_dir / "alpha" / "consistency-report.md").exists()
        assert (output_dir / "beta" / "grounding-report.md").exists()


class TestSemanticOnnx:
    """ONNX Runtime embedding path - torch-free.

    The pooling/normalisation unit tests inject a fake ONNX session +
    tokenizer so they run without any model download (and without torch).
    The real-model test is network-gated and skips when the e5 ONNX weights
    are not reachable.
    """

    def _make_grounder(self, fake_session, fake_tokenizer, input_names):
        from stellars_claude_code_plugins.document_processing.semantic import (
            SemanticGrounder,
        )

        g = SemanticGrounder.__new__(SemanticGrounder)
        g._session = fake_session
        g._tokenizer = fake_tokenizer
        g._input_names = set(input_names)
        g.model_name = "intfloat/multilingual-e5-small"
        return g

    def test_embed_mean_pool_unit_norm_and_token_type_ids(self):
        """_embed: mask-aware mean-pool, L2-norm, and zero token_type_ids feed."""
        import numpy as np

        def fake_tok(texts, **kw):
            n = len(texts)
            ids = np.array([[1, 2, 3], [4, 5, 0]], dtype="int64")[:n]
            mask = np.array([[1, 1, 1], [1, 1, 0]], dtype="int64")[:n]
            return {"input_ids": ids, "attention_mask": mask}

        captured: dict = {}

        class FakeSession:
            def run(self, _outputs, feed):
                captured.update(feed)
                # (N, L, dim=2). Text1 token2 has a huge value that MUST be
                # masked out by the attention_mask above.
                last = np.array(
                    [
                        [[1.0, 0.0], [3.0, 0.0], [5.0, 0.0]],
                        [[2.0, 2.0], [4.0, 4.0], [9.9, 9.9]],
                    ],
                    dtype="float32",
                )
                return [last]

        g = self._make_grounder(
            FakeSession(), fake_tok, {"input_ids", "attention_mask", "token_type_ids"}
        )
        vecs = g._embed(["a", "b"])

        # token_type_ids supplied as zeros, shaped like input_ids
        assert "token_type_ids" in captured
        assert captured["token_type_ids"].shape == (2, 3)
        assert int(captured["token_type_ids"].sum()) == 0

        # rows are unit-norm
        assert np.allclose(np.linalg.norm(vecs, axis=1), 1.0, atol=1e-6)

        # text0: mean (1,3,5)/3 = (3,0) -> normalised (1,0)
        assert np.allclose(vecs[0], [1.0, 0.0], atol=1e-6)
        # text1: masked token2 excluded -> mean (2,4)/2 = (3,3) -> (0.707,0.707)
        assert np.allclose(vecs[1], [2**-0.5, 2**-0.5], atol=1e-6)

    def test_cache_round_trip_skips_reembed(self, tmp_path):
        """_load_or_embed embeds once, then serves the parquet cache."""
        import numpy as np

        pytest.importorskip("pyarrow")
        from stellars_claude_code_plugins.document_processing.chunking import (
            recursive_chunk,
        )
        from stellars_claude_code_plugins.document_processing.semantic import (
            SemanticGrounder,
        )

        g = SemanticGrounder.__new__(SemanticGrounder)
        g.model_name = "intfloat/multilingual-e5-small"
        g.cache_dir = tmp_path
        g.max_chars = 1500
        calls = {"n": 0}

        def fake_embed(texts):
            calls["n"] += 1
            return np.full((len(texts), 4), 0.5, dtype="float32")

        g._embed = fake_embed

        text = "The estate has three walled gardens and an orchard. " * 20
        chunks = recursive_chunk(text, max_chars=1500)
        v1 = g._load_or_embed("doc.txt", text, chunks)
        v2 = g._load_or_embed("doc.txt", text, chunks)

        assert calls["n"] == 1  # second call hit the cache, no re-embed
        assert np.allclose(v1, v2)

    def test_real_model_paraphrase_outscores_unrelated(self, tmp_path):
        """Network-gated: real e5 ONNX ranks a paraphrase above an unrelated line."""
        pytest.importorskip("onnxruntime")
        pytest.importorskip("transformers")
        from stellars_claude_code_plugins.document_processing.semantic import (
            SemanticGrounder,
            is_available,
        )

        if not is_available():
            pytest.skip("semantic extras not installed")
        try:
            g = SemanticGrounder(cache_dir=str(tmp_path / "cache"))
        except Exception as exc:  # noqa: BLE001 - skip on no network / missing onnx weights
            pytest.skip(f"e5 ONNX model unavailable: {exc}")

        q = g._embed(["query: number of walled gardens on the estate"])[0]
        p_rel = g._embed(["passage: The estate has three walled gardens."])[0]
        p_unrel = g._embed(["passage: Rainfall averages 800 mm per year."])[0]

        assert float(q @ p_rel) > float(q @ p_unrel)
        assert g._is_e5() is True  # e5 query/passage prefixes apply
        assert 0.0 <= g.self_score("number of walled gardens") <= 1.0


class TestYearContradiction:
    """Regression for the contradiction-detection bug: historical years and
    years followed by a stopword must key on the 'year' category so a real
    contradiction (built 1650 vs source 1820) is caught, while a supported
    year (1998 present) is not flagged."""

    def test_historical_years_tagged_year(self):
        from stellars_claude_code_plugins.document_processing.entity_check import extract_numbers

        # 16xx/18xx (outside 19xx/20xx) must still be recognised as years
        assert ("1650", "", "year") in extract_numbers("built in 1650")
        # a year followed by the stopword "and" must not key on "and"
        nums = dict((v, cw) for v, _u, cw in extract_numbers("built in 1820 and restored in 1998"))
        assert nums.get("1820") == "year"
        assert nums.get("1998") == "year"

    def test_year_contradiction_caught(self):
        from stellars_claude_code_plugins.document_processing.entity_check import (
            find_numeric_mismatches,
        )

        passage = "The manor was built in 1820 and restored in 1998."
        assert find_numeric_mismatches("the manor was built in 1650", passage)  # contradiction
        assert not find_numeric_mismatches("the manor was restored in 1998", passage)  # supported

    def test_year_contradiction_through_ground(self):
        from stellars_claude_code_plugins.document_processing.grounding import ground

        src = [("e.txt", "The manor was built in 1820 and restored in 1998.")]
        m = ground("the manor was built in 1650", src)
        assert m.match_type == "contradicted"


class TestBm25IdfRecall:
    """Regression for the bm25 common-word false positive: a claim whose
    DISTINCTIVE tokens are absent must not confirm just because it shares
    ubiquitous words with the source (IDF-weighted recall)."""

    def test_common_word_fabrication_not_confirmed(self):
        from stellars_claude_code_plugins.document_processing.grounding import ground

        src = [
            (
                "e.txt",
                "The estate has three walled gardens.\n\n"
                "A trout stream runs along the eastern boundary.",
            )
        ]
        # shares only common tokens (estate, runs); distinctive (commercial,
        # brewery) absent -> must NOT be a bm25 confirm
        m = ground("the estate runs a commercial brewery", src)
        assert m.match_type not in ("exact", "fuzzy", "bm25")

    def test_distinctive_claim_still_confirms(self):
        from stellars_claude_code_plugins.document_processing.grounding import ground

        src = [("e.txt", "The estate has three walled gardens and an orchard.")]
        m = ground("the estate has three walled gardens", src)
        assert m.match_type in ("exact", "fuzzy", "bm25")


class TestGroundingEndToEnd:
    """Honest end-to-end test of the deterministic grounding (no model): on a
    realistic monolingual set - grounded claims, off-topic fabrications, and
    numeric contradictions - it meets the precision/recall targets. The
    cross-lingual on-topic case is the documented embedding ceiling, out of
    scope for the deterministic engine."""

    SRC = [
        (
            "estate.txt",
            "The estate has three walled gardens and an orchard.\n\n"
            "Rainfall in the region averages 800 millimetres per year.\n\n"
            "The manor was built in 1820 and restored in 1998.\n\n"
            "The vineyard covers twelve hectares on the south slope.\n\n"
            "A trout stream runs along the eastern boundary.",
        )
    ]
    CLAIMS = [
        ("the estate has three walled gardens", 1),
        ("there is an orchard on the estate", 1),
        ("rainfall averages 800 millimetres per year", 1),
        ("the manor was restored in 1998", 1),
        ("a trout stream runs along the eastern boundary", 1),
        ("the vineyard covers twelve hectares on the south slope", 1),
        ("the estate has a helicopter landing pad", 0),
        ("the estate runs a commercial brewery", 0),
        ("a private airport serves the estate", 0),
        ("the manor was built in 1650", 0),
        ("the manor was restored in 2010", 0),
        ("rainfall averages 200 millimetres per year", 0),
    ]

    def test_precision_recall_targets(self):
        from stellars_claude_code_plugins.document_processing.grounding import ground_many

        matches = ground_many([c for c, _ in self.CLAIMS], self.SRC)
        tp = fp = fn = 0
        for (_c, lab), m in zip(self.CLAIMS, matches):
            confirmed = m.match_type in ("exact", "fuzzy", "bm25")
            if lab and confirmed:
                tp += 1
            elif not lab and confirmed:
                fp += 1
            elif lab and not confirmed:
                fn += 1
        precision = tp / (tp + fp) if (tp + fp) else 0.0
        recall = tp / (tp + fn) if (tp + fn) else 0.0
        assert precision >= 0.90, f"precision {precision}"
        assert recall >= 0.80, f"recall {recall}"


class _FakeNLI:
    """Stub NLI grounder with fixed scores - exercises the wiring, no model."""

    def __init__(self, scores: dict):
        self._scores = scores

    def scores(self, premise: str, hypothesis: str) -> dict:
        return dict(self._scores)


class TestNLIGrounding:
    """NLI verdict wiring into ground() (fast, no model download)."""

    SRC = [("e.txt", "The estate has three walled gardens and an orchard.")]

    def test_entailment_grounds(self):
        from stellars_claude_code_plugins.document_processing.grounding import ground

        fake = _FakeNLI({"entailment": 0.95, "neutral": 0.03, "contradiction": 0.02})
        m = ground("le domaine possede trois jardins clos", self.SRC, nli_grounder=fake)
        assert m.match_type in ("exact", "fuzzy", "bm25", "semantic")  # grounded
        assert m.nli_scores["entailment"] == 0.95

    def test_contradiction_flagged(self):
        from stellars_claude_code_plugins.document_processing.grounding import ground

        fake = _FakeNLI({"entailment": 0.02, "neutral": 0.03, "contradiction": 0.95})
        m = ground("the estate has no gardens at all", self.SRC, nli_grounder=fake)
        assert m.match_type == "contradicted"

    def test_neutral_unconfirmed(self):
        from stellars_claude_code_plugins.document_processing.grounding import ground

        fake = _FakeNLI({"entailment": 0.10, "neutral": 0.80, "contradiction": 0.10})
        m = ground("qz zztop kvqj wbrtz", self.SRC, nli_grounder=fake)
        assert m.match_type == "none"

    def test_lexical_default_unaffected_without_nli(self):
        # No nli_grounder -> deterministic behaviour unchanged.
        from stellars_claude_code_plugins.document_processing.grounding import ground

        m = ground("the estate has three walled gardens", self.SRC)
        assert m.match_type == "exact"
        assert m.nli_scores == {}

    def test_extract_features_includes_nli(self):
        from stellars_claude_code_plugins.document_processing.grounding import (
            GroundingMatch,
            extract_features,
        )

        feat = extract_features(
            GroundingMatch(claim="x"),
            nli_scores={"entailment": 0.7, "neutral": 0.1, "contradiction": 0.2},
        )
        assert feat["nli_entail"] == 0.7
        assert feat["nli_contra"] == 0.2

    def test_real_model_entailment_and_crosslingual(self):
        pytest.importorskip("onnxruntime")
        pytest.importorskip("transformers")
        from stellars_claude_code_plugins.document_processing.nli import NLIGrounder, is_available

        if not is_available():
            pytest.skip("NLI extras not installed")
        try:
            g = NLIGrounder()
        except Exception as exc:  # noqa: BLE001 - skip on no network / missing weights
            pytest.skip(f"NLI model unavailable: {exc}")

        ev = "The estate has three walled gardens and an orchard."
        assert g.scores(ev, "There are three gardens on the estate.")["entailment"] > 0.5
        # cross-lingual entailment - the case cosine similarity could not solve
        assert g.scores(ev, "le domaine possede trois jardins clos")["entailment"] > 0.5
        assert g.verdict("The vineyard covers twelve hectares.", "The vineyard covers forty hectares.") == "contradicted"
